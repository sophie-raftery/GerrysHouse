"""
make_guide.py — generates How_To_Make_Gerrys_House.docx
Run: python make_guide.py  (from the GerrysHouse root folder)
"""
import os
from io import BytesIO
from PIL import Image, ImageDraw, ImageFont
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

doc = Document()

# ── page margins ──────────────────────────────────────────────────────────
for section in doc.sections:
    section.top_margin    = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin   = Inches(1.1)
    section.right_margin  = Inches(1.1)

# ── helper: add a coloured code block ────────────────────────────────────
def add_code(text, bg=(30, 30, 30)):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    shading = OxmlElement('w:shd')
    shading.set(qn('w:val'), 'clear')
    shading.set(qn('w:color'), 'auto')
    hex_bg = '{:02X}{:02X}{:02X}'.format(*bg)
    shading.set(qn('w:fill'), hex_bg)
    p._p.get_or_add_pPr().append(shading)
    run = p.add_run(text)
    run.font.name  = 'Courier New'
    run.font.size  = Pt(9)
    run.font.color.rgb = RGBColor(220, 220, 170)
    return p

def add_h1(text):
    p = doc.add_heading(text, level=1)
    p.runs[0].font.color.rgb = RGBColor(30, 100, 200)
    return p

def add_h2(text):
    p = doc.add_heading(text, level=2)
    p.runs[0].font.color.rgb = RGBColor(60, 140, 60)
    return p

def add_h3(text):
    return doc.add_heading(text, level=3)

def add_body(text):
    p = doc.add_paragraph(text)
    p.runs[0].font.size = Pt(11)
    return p

def add_img_from_pil(pil_img, width_inches=5.5, caption=""):
    buf = BytesIO()
    pil_img.save(buf, format='PNG')
    buf.seek(0)
    doc.add_picture(buf, width=Inches(width_inches))
    if caption:
        cp = doc.add_paragraph(caption)
        cp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        cp.runs[0].font.size = Pt(9)
        cp.runs[0].font.italic = True

def add_img_file(path, width_inches=5.5, caption=""):
    if os.path.isfile(path):
        doc.add_picture(path, width=Inches(width_inches))
        if caption:
            cp = doc.add_paragraph(caption)
            cp.alignment = WD_ALIGN_PARAGRAPH.CENTER
            cp.runs[0].font.size = Pt(9)
            cp.runs[0].font.italic = True

# ── diagram helpers ───────────────────────────────────────────────────────
def make_game_loop_diagram():
    W, H = 700, 220
    img = Image.new('RGB', (W, H), (245, 248, 255))
    d = ImageDraw.Draw(img)
    boxes = [
        (60,  80, 200, 140, (70, 130, 200),  "1. Handle\nEvents"),
        (270, 80, 430, 140, (60, 160, 80),   "2. Update\nGame State"),
        (480, 80, 640, 140, (200, 100, 50),  "3. Draw\nEverything"),
    ]
    for x1,y1,x2,y2,col,label in boxes:
        d.rounded_rectangle([x1,y1,x2,y2], radius=10, fill=col)
        d.multiline_text(((x1+x2)//2, (y1+y2)//2), label,
                         fill='white', anchor='mm', align='center')
    for i in range(len(boxes)-1):
        x1,y1,x2,y2,_,_ = boxes[i]
        nx1,ny1,nx2,ny2,_,_ = boxes[i+1]
        d.line([(x2, (y1+y2)//2), (nx1, (ny1+ny2)//2)], fill=(80,80,80), width=3)
        d.polygon([(nx1,(ny1+ny2)//2-6),(nx1+12,(ny1+ny2)//2),(nx1,(ny1+ny2)//2+6)],
                  fill=(80,80,80))
    # loop back arrow
    d.line([(620,140),(620,180),(80,180),(80,140)], fill=(150,50,150), width=2)
    d.polygon([(74,140),(80,128),(86,140)], fill=(150,50,150))
    d.text((340,195), "Repeats 60 times per second", fill=(100,100,100), anchor='mm')
    d.text((350,30), "The Game Loop", fill=(30,30,30), anchor='mm')
    return img

def make_file_structure_diagram():
    W, H = 620, 340
    img = Image.new('RGB', (W, H), (250, 250, 250))
    d = ImageDraw.Draw(img)
    lines = [
        (0,  "GerrysHouse/",            (30,30,30),   True),
        (1,  "├── Level1/",             (60,100,180), True),
        (2,  "│   ├── Test_level.py",   (80,80,80),   False),
        (2,  "│   ├── Lvl 1 Gerry room.py", (80,80,80), False),
        (2,  "│   ├── door.py",         (180,100,30), False),
        (2,  "│   ├── hotbar.py",       (180,100,30), False),
        (2,  "│   └── shared_state.py", (180,100,30), False),
        (1,  "├── Level 2/",            (60,160,60),  True),
        (2,  "│   ├── Lvl 2.py",        (80,80,80),   False),
        (2,  "│   └── Kitchen Lvl 2.py",(80,80,80),   False),
        (1,  "├── Level 3/",            (160,80,160), True),
        (2,  "│   ├── Lvl 3.py",        (80,80,80),   False),
        (2,  "│   └── Garage Lvl 3.py", (80,80,80),   False),
        (1,  "├── Level 4/",            (200,120,30), True),
        (1,  "├── Vivienne's room/",    (100,160,160),True),
        (1,  "└── images/",             (120,120,120),True),
    ]
    y = 15
    for indent, text, color, bold in lines:
        d.text((15 + indent*20, y), text, fill=color)
        y += 19
    return img

def make_collision_diagram():
    W, H = 600, 320
    img = Image.new('RGB', (W, H), (30, 40, 30))
    d = ImageDraw.Draw(img)
    # walls
    d.rectangle([0,0,W,30],  fill=(180,60,60), outline=(255,80,80))
    d.rectangle([0,290,W,H], fill=(180,60,60), outline=(255,80,80))
    d.rectangle([0,0,30,H],  fill=(180,60,60), outline=(255,80,80))
    d.rectangle([570,0,W,H], fill=(180,60,60), outline=(255,80,80))
    # furniture block
    d.rectangle([200,100,380,200], fill=(120,80,40), outline=(200,140,60))
    d.text((290,145), "Collision Rect\n(furniture)", fill='white', anchor='mm', align='center')
    # player
    d.ellipse([280,220,320,260], fill=(80,160,240))
    d.text((300,240), "P", fill='white', anchor='mm')
    # arrows showing push-out
    d.line([(300,215),(300,205)], fill=(255,255,100), width=2)
    d.polygon([(296,205),(300,195),(304,205)], fill=(255,255,100))
    d.text((300,175), "Player pushed OUT", fill=(255,255,100), anchor='mm')
    d.text((290, 60), "Red = Collision Rectangles (invisible in game)",
           fill=(255,150,150), anchor='mm')
    return img

def make_agro_diagram():
    W, H = 480, 300
    img = Image.new('RGB', (W, H), (20, 30, 20))
    d = ImageDraw.Draw(img)
    cx, cy = 240, 150
    # agro ring
    for r in range(140, 145):
        d.ellipse([cx-r, cy-r, cx+r, cy+r], outline=(200,50,50,180))
    # fill
    d.ellipse([cx-140,cy-140,cx+140,cy+140], fill=(60,20,20), outline=(200,50,50))
    d.text((cx, cy-80), "Agro Zone", fill=(255,100,100), anchor='mm')
    # npc
    d.ellipse([cx-16,cy-16,cx+16,cy+16], fill=(200,80,80))
    d.text((cx,cy), "NPC", fill='white', anchor='mm')
    # player outside
    px, py = 390, 150
    d.ellipse([px-14,py-14,px+14,py+14], fill=(80,160,240))
    d.text((px,py), "P", fill='white', anchor='mm')
    d.text((px, py+25), "Outside — safe", fill=(100,200,100), anchor='mm')
    # player inside
    px2, py2 = 160, 200
    d.ellipse([px2-14,py2-14,px2+14,py2+14], fill=(80,160,240))
    d.text((px2,py2), "P", fill='white', anchor='mm')
    d.text((px2, py2+25), "Inside — CHASED!", fill=(255,100,100), anchor='mm')
    # chase arrow
    d.line([(cx-10,cy+10),(px2+14,py2-5)], fill=(255,200,50), width=2)
    return img

def make_ysort_diagram():
    W, H = 560, 280
    img = Image.new('RGB', (W, H), (40, 55, 40))
    d = ImageDraw.Draw(img)
    # counter
    d.rectangle([80, 80, 480, 200], fill=(100, 70, 40), outline=(200, 150, 80), width=2)
    d.text((280, 140), "Kitchen Counter / Island", fill=(220,180,100), anchor='mm')
    # player behind (above centre)
    d.ellipse([220, 90, 260, 130], fill=(80,160,240))
    d.text((240,110), "P", fill='white', anchor='mm')
    d.text((240, 70), "Player above centre → drawn FIRST\n(counter appears on top)",
           fill=(200,200,255), anchor='mm', align='center')
    # player in front (below centre)
    d.ellipse([310, 195, 350, 235], fill=(80,160,240))
    d.text((330,215), "P", fill='white', anchor='mm')
    d.text((330, 255), "Player below centre → drawn LAST\n(player appears on top)",
           fill=(200,255,200), anchor='mm', align='center')
    return img

def make_level_flow_diagram():
    W, H = 660, 420
    img = Image.new('RGB', (W, H), (248, 248, 255))
    d = ImageDraw.Draw(img)
    nodes = [
        (330, 30,  "Test Level\n(L1 Garden)",   (60,120,200)),
        (330, 110, "Gerry's Room L1",            (80,160,80)),
        (330, 190, "L2 Garden",                  (60,120,200)),
        (330, 270, "Kitchen L2",                 (80,160,80)),
        (330, 350, "L3 Garden → Garage → L4 → Final Room",
                                                 (160,80,160)),
    ]
    loses = [
        (100, 110, "LOSE\nScreen",  (200,60,60)),
        (100, 270, "LOSE\nScreen",  (200,60,60)),
    ]
    for lx,ly,lt,lc in loses:
        d.rounded_rectangle([lx-45,ly-25,lx+45,ly+25], radius=8, fill=lc)
        d.multiline_text((lx,ly), lt, fill='white', anchor='mm', align='center')
    for nx,ny,nt,nc in nodes:
        tw = max(160, len(nt)*7)
        d.rounded_rectangle([nx-tw//2,ny-28,nx+tw//2,ny+28], radius=8, fill=nc)
        d.multiline_text((nx,ny), nt, fill='white', anchor='mm', align='center')
    for i in range(len(nodes)-1):
        nx,ny,_,_ = nodes[i]
        nnx,nny,_,_ = nodes[i+1]
        d.line([(nx,ny+28),(nnx,nny-28)], fill=(80,80,80), width=2)
        d.polygon([(nnx-5,nny-28),(nnx,nny-18),(nnx+5,nny-28)], fill=(80,80,80))
    # lose arrows
    d.line([(100,85),(100,110-25)], fill=(200,60,60), width=1)
    d.text((60, 70), "Dog catches\nplayer", fill=(200,60,60), anchor='mm', align='center')
    d.line([(100,245),(100,270-25)], fill=(200,60,60), width=1)
    d.text((60, 230), "Mother\ncatches", fill=(200,60,60), anchor='mm', align='center')
    return img

def make_hotbar_diagram():
    W, H = 560, 110
    img = Image.new('RGB', (W, H), (30, 30, 30))
    d = ImageDraw.Draw(img)
    slot_size = 70
    start_x = (W - 5 * slot_size - 4 * 8) // 2
    items = ["Shovel", "Dog\nBone", "Room\nKey", "MJ\nVinyl", ""]
    for i, item in enumerate(items):
        x = start_x + i * (slot_size + 8)
        color = (80, 120, 200) if i == 0 else (50, 50, 50)
        border = (255, 220, 50) if i == 0 else (150, 150, 150)
        d.rounded_rectangle([x, 15, x+slot_size, 15+slot_size],
                             radius=6, fill=color, outline=border, width=3 if i==0 else 1)
        if item:
            d.multiline_text((x+slot_size//2, 15+slot_size//2), item,
                             fill='white', anchor='mm', align='center')
        d.text((x+slot_size//2, 95), str(i+1), fill=(200,200,200), anchor='mm')
    d.text((W//2, 5), "Hotbar — press 1-5 to select a slot", fill=(200,200,200), anchor='mm')
    return img


# ═══════════════════════════════════════════════════════════════════════════
# DOCUMENT CONTENT
# ═══════════════════════════════════════════════════════════════════════════

# ── Title page ───────────────────────────────────────────────────────────
title = doc.add_heading("How to Make Gerry's House", 0)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
title.runs[0].font.color.rgb = RGBColor(30, 80, 180)
title.runs[0].font.size = Pt(28)

sub = doc.add_paragraph("A Complete Python Game Development Guide for Ages 16+")
sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
sub.runs[0].font.size = Pt(14)
sub.runs[0].font.color.rgb = RGBColor(100,100,100)
sub.runs[0].font.italic = True

doc.add_paragraph("")

# splash — use garden background if available
if os.path.isfile("images/garden.png"):
    add_img_file("images/garden.png", width_inches=5.5,
                 caption="Level 1 Garden — the starting area of the game")

doc.add_page_break()

# ── Contents ─────────────────────────────────────────────────────────────
add_h1("Table of Contents")
toc_items = [
    "1. What Is This Game?",
    "2. What You Need (Setup)",
    "3. How Python Games Work — The Game Loop",
    "4. File Structure Explained",
    "5. Making the Player Move",
    "6. Collision Detection",
    "7. The Hotbar (Inventory System)",
    "8. Interactable Objects",
    "9. The Door System — Moving Between Levels",
    "10. Passing Inventory Between Levels",
    "11. NPC Patrol AI and the Agro Ring",
    "12. Y-Sorting — Walking Behind Objects",
    "13. The Full Level Flow",
    "14. Every File Explained",
    "15. Adding a New Level (Template)",
    "16. Common Mistakes and Fixes",
]
for item in toc_items:
    p = doc.add_paragraph(item, style='List Bullet')
    p.runs[0].font.size = Pt(11)

doc.add_page_break()

# ═══ SECTION 1 ════════════════════════════════════════════════════════════
add_h1("1. What Is This Game?")
add_body(
    "Gerry's House is a 2D top-down adventure game built entirely in Python using a library "
    "called pygame-ce. The player controls a character who must sneak around their house "
    "completing tasks — digging up a bone to distract a dog, making Gerry's bed, collecting "
    "three vinyl records, and finally playing them on a record player to trigger the ending."
)
add_body(
    "There is no game engine. Every single mechanic — movement, collision, animation, "
    "inventory, AI, transitions — is written by hand in Python. This makes it a brilliant "
    "learning project because you can see exactly how everything works."
)
doc.add_paragraph("")

# player sprite strip
sprites = []
for i in range(1, 5):
    p = f"images/Player_sprites/sprite-1-{i} (1).png"
    if os.path.isfile(p):
        sprites.append(Image.open(p).convert("RGBA"))

if sprites:
    max_h = max(s.height for s in sprites)
    strip = Image.new("RGBA", (sum(s.width for s in sprites) + 8*(len(sprites)-1), max_h),
                      (200, 200, 200, 0))
    x = 0
    for s in sprites:
        strip.paste(s, (x, 0), s)
        x += s.width + 8
    strip = strip.resize((min(400, strip.width * 3), min(120, max_h * 3)),
                         Image.NEAREST)
    add_img_from_pil(strip, width_inches=3.5,
                     caption="The player's walk-forward animation — 4 frames")

doc.add_page_break()

# ═══ SECTION 2 ════════════════════════════════════════════════════════════
add_h1("2. What You Need (Setup)")
add_h2("Install Python")
add_body("Download Python 3.10 or newer from python.org. During installation tick the box that says 'Add Python to PATH'.")

add_h2("Install pygame-ce")
add_body("Open a terminal (press Win+R, type cmd, press Enter) and type:")
add_code("pip install pygame-ce")

add_h2("Install a Code Editor")
add_body("VS Code is the best free option. Download it from code.visualstudio.com. Install the Python extension inside VS Code.")

add_h2("Run the Game")
add_body("Open a terminal in the GerrysHouse folder and type:")
add_code("python Level1/Test_level.py")

add_h2("Controls")
tbl = doc.add_table(rows=5, cols=2)
tbl.style = 'Table Grid'
rows_data = [("Key","Action"),("W A S D","Move player"),
             ("Shift","Sprint"),("E","Interact with objects"),("1-5","Select hotbar slot")]
for i,(k,v) in enumerate(rows_data):
    tbl.rows[i].cells[0].text = k
    tbl.rows[i].cells[1].text = v
    if i == 0:
        for cell in tbl.rows[i].cells:
            cell.paragraphs[0].runs[0].font.bold = True

doc.add_page_break()


# ═══ SECTION 3 ════════════════════════════════════════════════════════════
add_h1("3. How Python Games Work — The Game Loop")
add_body(
    "Every pygame game runs the same three-step loop over and over, typically 60 times "
    "per second. This is called the game loop. Without it, nothing in the game would move "
    "or respond to input."
)
add_img_from_pil(make_game_loop_diagram(), width_inches=5.5,
                 caption="The three steps of the game loop, running 60 times per second")

add_h2("Step 1 — Handle Events")
add_body(
    "pygame collects things that happen — key presses, mouse clicks, window close. "
    "You loop through them and respond. If you don't handle QUIT, the window won't close."
)
add_code(
    "for event in pygame.event.get():\n"
    "    if event.type == pygame.QUIT:\n"
    "        running = False\n"
    "    if event.type == pygame.KEYDOWN:\n"
    "        if event.key == pygame.K_e:\n"
    "            print('Player pressed E')"
)

add_h2("Step 2 — Update Game State")
add_body(
    "This is where everything moves. The player's position updates, timers tick, "
    "collisions are resolved, NPC AI runs. The variable 'dt' is the time in seconds "
    "since the last frame. Multiplying movement by dt makes speed consistent regardless "
    "of frame rate."
)
add_code(
    "dt = clock.tick(60) / 1000   # aim for 60fps, dt is ~0.016 seconds\n"
    "player.rect.x += speed * dt  # moves 150px per second at any frame rate"
)

add_h2("Step 3 — Draw Everything")
add_body(
    "Draw the background first (it covers everything), then sprites, then the HUD on top. "
    "Order matters — things drawn later appear on top."
)
add_code(
    "display_surface.blit(background, (0, 0))  # background first\n"
    "all_sprites.draw(display_surface)          # then sprites\n"
    "overlay.display(display_surface)           # HUD on top\n"
    "pygame.display.update()                    # show the frame"
)
doc.add_page_break()

# ═══ SECTION 4 ════════════════════════════════════════════════════════════
add_h1("4. File Structure Explained")
add_body(
    "The project is split into many files so each level is self-contained. Shared systems "
    "like the inventory and door transitions live in Level1/ and get imported by every other level."
)
add_img_from_pil(make_file_structure_diagram(), width_inches=4.5,
                 caption="Project file structure — orange files are shared by all levels")

add_h2("Shared Files (in Level1/)")
tbl = doc.add_table(rows=6, cols=2)
tbl.style = 'Table Grid'
for k,v in [("File","What it does"),
            ("door.py","Door class — proximity detection, fade, level loading"),
            ("hotbar.py","Inventory — InventoryItem, Hotbar, Overlay HUD"),
            ("shared_state.py","3 global variables that pass data between levels"),
            ("minesweeper.py","The bed mini-game launched in Gerry's room"),
            ("Test_level.py","The starting level (Level 1 garden)")]:
    row = tbl.add_row() if k != "File" else tbl.rows[0]
    i = ["File","door.py","hotbar.py","shared_state.py","minesweeper.py","Test_level.py"].index(k)
    tbl.rows[i].cells[0].text = k
    tbl.rows[i].cells[1].text = v
    if i == 0:
        for cell in tbl.rows[i].cells:
            cell.paragraphs[0].runs[0].font.bold = True

doc.add_page_break()

# ═══ SECTION 5 ════════════════════════════════════════════════════════════
add_h1("5. Making the Player Move")
add_body(
    "The player is a pygame Sprite — a class that bundles an image and a rect (rectangle "
    "position) together. pygame.sprite.Group manages all sprites and can draw them all at once."
)
add_code(
    "class Player(pygame.sprite.Sprite):\n"
    "    def __init__(self, groups):\n"
    "        super().__init__(groups)           # register with the group\n"
    "        self.image = pygame.image.load('images/Player_sprites/sprite-1-1 (1).png')\n"
    "        self.rect  = self.image.get_frect(center=(640, 360))  # position\n"
    "        self.direction = pygame.Vector2()  # x and y movement direction\n"
    "        self.speed = 150                   # pixels per second\n\n"
    "    def update(self, dt):\n"
    "        keys = pygame.key.get_pressed()    # which keys are held right now?\n"
    "        self.direction.x = int(keys[pygame.K_d]) - int(keys[pygame.K_a])  # right - left\n"
    "        self.direction.y = int(keys[pygame.K_s]) - int(keys[pygame.K_w])  # down - up\n"
    "        if self.direction.length() > 0:\n"
    "            self.direction = self.direction.normalize()  # same speed diagonally\n"
    "        self.rect.center += self.direction * self.speed * dt"
)
add_body(
    "The normalize() call is important. Without it, moving diagonally would be ~41% faster "
    "because the vector (1,1) has length 1.41, not 1."
)

add_h2("Animation")
add_body(
    "The game has 8 walking directions, each with 4 frames of animation. They're loaded as "
    "lists at the start of each level. Every 200ms the frame index advances by 1."
)
add_code(
    "walk_forward = [pygame.image.load(f'sprite-1-{i}.png') for i in range(1,5)]\n"
    "# frame_idx cycles 0 → 1 → 2 → 3 → 0 → 1 ...\n"
    "if now - last_updated > 200:\n"
    "    last_updated = now\n"
    "    frame_idx = (frame_idx + 1) % 4  # wrap back to 0 after 3\n"
    "self.image = walk_forward[frame_idx]"
)

if os.path.isfile("images/Gerry's room.png"):
    add_img_file("images/Gerry's room.png", width_inches=5.0,
                 caption="Gerry's Room — the first indoor level the player enters")

doc.add_page_break()


# ═══ SECTION 6 ════════════════════════════════════════════════════════════
add_h1("6. Collision Detection")
add_body(
    "Collision detection stops the player walking through walls and furniture. The game uses "
    "axis-aligned bounding box (AABB) collision with a list of invisible rectangles."
)
add_img_from_pil(make_collision_diagram(), width_inches=5.0,
                 caption="Collision rectangles are invisible in the game but block the player's movement")

add_h2("How It Works")
add_body("Each frame, after moving the player, call resolve_collision(player). It loops through every collision rectangle and checks for overlap.")
add_code(
    "COLLISION_RECTS = [\n"
    "    pygame.Rect(0,   0,    1280, 40),   # top wall\n"
    "    pygame.Rect(0,   680,  1280, 40),   # bottom wall\n"
    "    pygame.Rect(0,   0,    40,  720),   # left wall\n"
    "    pygame.Rect(1240,0,    40,  720),   # right wall\n"
    "    pygame.Rect(200, 100,  180, 100),   # a piece of furniture\n"
    "]\n\n"
    "def resolve_collision(player):\n"
    "    for wall in COLLISION_RECTS:\n"
    "        if not player.rect.colliderect(wall):\n"
    "            continue                    # not overlapping — skip\n"
    "        # How much overlap is there in each direction?\n"
    "        ox = min(player.rect.right - wall.left, wall.right - player.rect.left)\n"
    "        oy = min(player.rect.bottom - wall.top, wall.bottom - player.rect.top)\n"
    "        # Push out by the smaller overlap (less jarring)\n"
    "        if ox < oy:\n"
    "            if player.rect.centerx < wall.centerx:\n"
    "                player.rect.right = wall.left   # push left\n"
    "            else:\n"
    "                player.rect.left = wall.right   # push right\n"
    "        else:\n"
    "            if player.rect.centery < wall.centery:\n"
    "                player.rect.bottom = wall.top   # push up\n"
    "            else:\n"
    "                player.rect.top = wall.bottom   # push down"
)
add_body("Call this after all_sprites.update(dt) in the game loop. NPCs use the same function — just pass the NPC sprite instead of the player.")
doc.add_page_break()

# ═══ SECTION 7 ════════════════════════════════════════════════════════════
add_h1("7. The Hotbar (Inventory System)")
add_body(
    "The hotbar is a 5-slot inventory displayed at the bottom of the screen. It's defined "
    "in hotbar.py and used identically across every level."
)
add_img_from_pil(make_hotbar_diagram(), width_inches=5.0,
                 caption="The 5-slot hotbar — slot 1 is selected (yellow border)")

add_h2("InventoryItem")
add_body("An item is just a name, a category, and an image path.")
add_code(
    "from hotbar import InventoryItem\n\n"
    "room_key    = InventoryItem('Room_Key',    'Key Item',   'images/Key.png')\n"
    "mj_vinyl    = InventoryItem('MJ_Vinyl',    'Quest Item', 'images/items/Vinyl_white.png')\n"
    "dog_bone    = InventoryItem('Dog_Bone',    'Quest Item', 'images/items/Dog_Bone.png')"
)

add_h2("Adding and Checking Items")
add_code(
    "hotbar.add_item_first_free(room_key)     # add to first empty slot\n\n"
    "# Check if player has a specific item:\n"
    "has_key = any(s and s.name == 'Room_Key' for s in hotbar.slots)\n\n"
    "# Remove an item:\n"
    "for i, slot in enumerate(hotbar.slots):\n"
    "    if slot and slot.name == 'Room_Key':\n"
    "        hotbar.slots[i] = None\n"
    "        break"
)

add_h2("The Overlay (Drawing the Hotbar)")
add_code(
    "overlay = Overlay(Player)    # create once before the game loop\n\n"
    "# Inside the game loop, handle key selection:\n"
    "overlay.hotbar.handle_keypress(event)\n\n"
    "# At the end of the draw section:\n"
    "overlay.display(display_surface)"
)
doc.add_page_break()

# ═══ SECTION 8 ════════════════════════════════════════════════════════════
add_h1("8. Interactable Objects")
add_body(
    "Every interactable in the game — doors, boxes, dog bowls, dirt mounds, vinyl players — "
    "follows the same pattern. This consistency makes adding new interactables quick."
)

add_h2("The Pattern")
tbl = doc.add_table(rows=5, cols=2)
tbl.style = 'Table Grid'
for k,v in [("Part","Purpose"),
            ("pos / rect","Where the object is in the world"),
            ("INTERACT_RADIUS","How close the player must be for the prompt to show"),
            ("show_prompt","True/False flag — turns on glow and text when player is nearby"),
            ("draw()","Draws the glow ellipse and [E] prompt text when show_prompt is True")]:
    i = ["Part","pos / rect","INTERACT_RADIUS","show_prompt","draw()"].index(k)
    tbl.rows[i].cells[0].text = k
    tbl.rows[i].cells[1].text = v
    if i == 0:
        for cell in tbl.rows[i].cells:
            cell.paragraphs[0].runs[0].font.bold = True

doc.add_paragraph("")
add_code(
    "class KeyBox:\n"
    "    INTERACT_RADIUS = 90\n\n"
    "    def __init__(self, pos):\n"
    "        self.pos = pygame.Vector2(pos)\n"
    "        self.collected = False\n"
    "        self.show_prompt = False\n"
    "        self.rect = pygame.Rect(0, 0, 48, 48)\n"
    "        self.rect.center = (int(pos[0]), int(pos[1]))\n"
    "        font = pygame.font.SysFont(None, 20)\n"
    "        self._prompt = font.render('[E] Open', True, (255, 255, 255))\n"
    "        self._shadow = font.render('[E] Open', True, (0, 0, 0))\n\n"
    "    def update(self, player):\n"
    "        dist = pygame.Vector2(player.rect.center).distance_to(self.pos)\n"
    "        self.show_prompt = dist <= self.INTERACT_RADIUS and not self.collected\n\n"
    "    def draw(self, surface):\n"
    "        if self.show_prompt:\n"
    "            glow = pygame.Surface((70, 70), pygame.SRCALPHA)\n"
    "            pygame.draw.ellipse(glow, (255, 220, 100, 70), glow.get_rect())\n"
    "            surface.blit(glow, glow.get_rect(center=self.rect.center))\n"
    "            px = self.rect.centerx - self._prompt.get_width() // 2\n"
    "            py = self.rect.top - 22\n"
    "            surface.blit(self._shadow, (px+1, py+1))   # shadow for readability\n"
    "            surface.blit(self._prompt, (px, py))\n\n"
    "# In the game loop, on keypress E:\n"
    "if event.key == pygame.K_e:\n"
    "    dist = pygame.Vector2(player.rect.center).distance_to(key_box.pos)\n"
    "    if dist <= KeyBox.INTERACT_RADIUS and not key_box.collected:\n"
    "        key_box.collected = True\n"
    "        overlay.hotbar.add_item_first_free(room_key)"
)
doc.add_page_break()


# ═══ SECTION 9 ════════════════════════════════════════════════════════════
add_h1("9. The Door System — Moving Between Levels")
add_body(
    "door.py contains the Door class which handles everything about transitioning between levels. "
    "Any .py file with a run() function can be a level."
)
add_code(
    "from door import Door\n\n"
    "front_door = Door(\n"
    "    pos           = (1004, 320),               # centre position on screen\n"
    "    target_module = 'Level1/Lvl 1 Gerry room.py',  # next level file\n"
    "    image_path    = None,                      # None = invisible door\n"
    "    size          = (40, 60),                  # hitbox size in pixels\n"
    "    interact_radius = 80,                      # how close to show prompt\n"
    ")\n\n"
    "# Every frame in update section:\n"
    "front_door.update(player)   # sets show_prompt when player is close\n\n"
    "# On keypress E:\n"
    "if front_door.try_enter(player):\n"
    "    front_door.transition(display_surface)     # black fade-out\n"
    "    front_door.load_next_level()               # run the next file\n\n"
    "# Every frame in draw section:\n"
    "front_door.draw(display_surface)              # shows [E] Enter when nearby"
)

add_h2("How load_next_level() Works")
add_body(
    "It uses Python's importlib to load any .py file dynamically and call its run() function. "
    "This means levels don't need to import each other — they just point to a file path."
)
add_code(
    "import importlib.util\n\n"
    "spec   = importlib.util.spec_from_file_location('_next_level', 'Level1/Lvl 1 Gerry room.py')\n"
    "module = importlib.util.module_from_spec(spec)\n"
    "spec.loader.exec_module(module)    # runs the file\n"
    "module.run()                       # calls run() inside it"
)
doc.add_page_break()

# ═══ SECTION 10 ═══════════════════════════════════════════════════════════
add_h1("10. Passing Inventory Between Levels")
add_body(
    "When the player moves to a new level, their hotbar items need to carry over. "
    "shared_state.py is a tiny module that acts as a bulletin board between levels."
)
add_code(
    "# shared_state.py — the whole file:\n"
    "incoming_hotbar_slots = None   # items going INTO the next level\n"
    "returned_hotbar_slots = None   # items coming BACK\n"
    "player_spawn          = None   # where to place the player on return"
)
add_body("When leaving a level, write the current hotbar to shared_state before loading the next level:")
add_code(
    "import shared_state\n"
    "shared_state.incoming_hotbar_slots = list(overlay.hotbar.slots)\n"
    "front_door.load_next_level()"
)
add_body("At the start of the next level's run(), read it back:")
add_code(
    "if getattr(shared_state, 'incoming_hotbar_slots', None):\n"
    "    for i, item in enumerate(shared_state.incoming_hotbar_slots):\n"
    "        overlay.hotbar.slots[i] = item\n"
    "    shared_state.incoming_hotbar_slots = None  # clear it"
)
doc.add_page_break()

# ═══ SECTION 11 ═══════════════════════════════════════════════════════════
add_h1("11. NPC Patrol AI and the Agro Ring")
add_body(
    "The dog, the mother in the kitchen, and the father in the garage all use the same AI. "
    "It has three modes: patrol, chase, and (for the father) pause."
)
add_img_from_pil(make_agro_diagram(), width_inches=4.5,
                 caption="The agro ring — once the player steps inside it, the NPC switches to chase mode")

add_h2("Patrol Waypoints")
add_body("The NPC loops through a list of positions. When it reaches one, it moves to the next.")
add_code(
    "PATROL = [\n"
    "    pygame.Vector2(300, 110),   # top-left\n"
    "    pygame.Vector2(680, 100),   # top-right\n"
    "    pygame.Vector2(730, 450),   # right side\n"
    "]\n\n"
    "def _next_patrol_point(self):\n"
    "    self._patrol_idx = (self._patrol_idx + 1) % len(self.PATROL)\n"
    "    self._waypoint   = pygame.Vector2(self.PATROL[self._patrol_idx])"
)

add_h2("Steering / Obstacle Avoidance")
add_body(
    "Before moving, the NPC fires three feeler rays ahead of it at -35°, 0°, and +35°. "
    "If a ray hits a collision rect, the NPC steers away. This stops it getting stuck in corners."
)
add_code(
    "FEELER_ANGLES = (-35, 0, 35)   # three rays: left, centre, right\n\n"
    "def _compute_steering(self, goal_dir):\n"
    "    avoidance = pygame.Vector2(0, 0)\n"
    "    for angle in self.FEELER_ANGLES:\n"
    "        avoidance += self._cast_feeler(self.pos, angle, goal_dir)\n"
    "    if avoidance.length() > 0:\n"
    "        steering = goal_dir + avoidance * 2.5\n"
    "        return steering.normalize()\n"
    "    return goal_dir"
)

add_h2("Agro Ring — Chase Mode")
add_code(
    "AGRO_RADIUS = 150   # pixels — about 3× the player's size\n\n"
    "def update(self, dt):\n"
    "    player_dist = pygame.Vector2(self._player.rect.center).distance_to(self.pos)\n"
    "    self._chasing = player_dist <= self.AGRO_RADIUS\n\n"
    "    if self._chasing:\n"
    "        to_player = pygame.Vector2(self._player.rect.center) - self.pos\n"
    "        move_dir  = to_player.normalize()\n"
    "        self.pos += move_dir * self.CHASE_SPEED * dt   # faster than patrol\n"
    "    else:\n"
    "        # normal patrol code here\n"
    "        ..."
)

add_h2("Catch Detection")
add_body("If the NPC touches the player while chasing, the game transitions to the losing screen.")
add_code(
    "# In the game loop, after all_sprites.update(dt):\n"
    "if npc._chasing and npc.rect.colliderect(player.rect):\n"
    "    lose_door.transition(display_surface)\n"
    "    lose_door.load_next_level()\n"
    "    return"
)
doc.add_page_break()

# ═══ SECTION 12 ═══════════════════════════════════════════════════════════
add_h1("12. Y-Sorting — Walking Behind Objects")
add_body(
    "In a top-down game, objects lower on screen should appear in front of objects higher up. "
    "The kitchen counter and garage car both use this technique."
)
add_img_from_pil(make_ysort_diagram(), width_inches=5.0,
                 caption="Y-sorting: player above the counter's midpoint appears behind it; below appears in front")
add_code(
    "# Split all sprites into two groups based on Y position\n"
    "behind  = [s for s in all_sprites if s.rect.bottom < counter_rect.centery]\n"
    "infront = [s for s in all_sprites if s.rect.bottom >= counter_rect.centery]\n\n"
    "# Draw in order: back sprites → counter → front sprites\n"
    "for sprite in behind:\n"
    "    display_surface.blit(sprite.image, sprite.rect)\n"
    "display_surface.blit(counter_img, counter_rect)   # the counter\n"
    "for sprite in infront:\n"
    "    display_surface.blit(sprite.image, sprite.rect)"
)
add_body("This applies to both the player and the Mother NPC simultaneously — any sprite above the midpoint is 'behind'.")

if os.path.isfile("images/kitchen.png"):
    add_img_file("images/kitchen.png", width_inches=5.0,
                 caption="The kitchen — the counter uses y-sorting so the player walks behind it")

doc.add_page_break()

# ═══ SECTION 13 ═══════════════════════════════════════════════════════════
add_h1("13. The Full Level Flow")
add_body("The game follows a fixed path through all levels. Here is the complete route from start to finish:")
add_img_from_pil(make_level_flow_diagram(), width_inches=5.5,
                 caption="The game's level progression — follow the arrows from top to bottom")

add_h2("What You Need at Each Stage")
tbl = doc.add_table(rows=6, cols=2)
tbl.style = 'Table Grid'
for k,v in [("Level","What to do"),
            ("L1 Garden","Dig up the bone with the shovel. Put it in the dog bowl to distract the dog. Enter Gerry's room."),
            ("Gerry's Room L1","Play minesweeper to make the bed → MJ_Vinyl. Open the key box → Room_Key. Exit with both."),
            ("L2 Garden → Kitchen","Enter the kitchen. Pick up the Billy_Vinyl from the counter. Avoid Mother. Exit."),
            ("L3 Garden → Garage","Enter the garage. Pick up the Katie_Vinyl from the vinyl box. Avoid Father. Exit."),
            ("L4 → Final Room","Place all 3 vinyls on the record player. Watch the cutscene. Win!")]:
    i = ["Level","L1 Garden","Gerry's Room L1","L2 Garden → Kitchen","L3 Garden → Garage","L4 → Final Room"].index(k)
    tbl.rows[i].cells[0].text = k
    tbl.rows[i].cells[1].text = v
    if i == 0:
        for cell in tbl.rows[i].cells:
            cell.paragraphs[0].runs[0].font.bold = True

doc.add_page_break()


# ═══ SECTION 14 ═══════════════════════════════════════════════════════════
add_h1("14. Every File Explained")

files_info = [
    ("shared_state.py", "A 3-variable module used as a data bus between levels. Contains incoming_hotbar_slots (items going into the next level), returned_hotbar_slots (items coming back), and player_spawn (spawn position on return)."),
    ("hotbar.py", "Defines InventoryItem (name + image), Hotbar (a 5-slot list), and Overlay (draws the hotbar HUD to the screen). Used identically in every single level."),
    ("door.py", "The Door class. Handles proximity detection (try_enter), the black fade transition, and dynamic loading of any .py level file. Making a door invisible is as simple as passing image_path=None."),
    ("minesweeper.py", "A minesweeper mini-game launched when the player interacts with Gerry's bed. Returns True if the player won, False if they lost."),
    ("Test_level.py", "Level 1 — the garden. Contains: GroundItem (pickups), DogBowl (place bone here to distract dog), DirtMound (dig with shovel to get bone), Dog (patrol + agro AI), ExitGate (vinyl check)."),
    ("Lvl 1 Gerry room.py", "Gerry's bedroom. InteractableBox gives Room_Key. MakeBed launches minesweeper and rewards MJ_Vinyl. Exit door requires both vinyl + key. Uses Door class to return to Test_level."),
    ("Lvl 2.py", "Level 2 garden. Single door to Kitchen Lvl 2 and a vinyl exit door. Loads Garage(out) image for the garage exterior."),
    ("Kitchen Lvl 2.py", "The kitchen interior. Mother NPC patrols using steering AI. AGRO_RADIUS=300px. Billy_Vinyl sits on the counter. Picking it up speeds Mother up. Y-sorting makes counter layer correctly. Lose door goes to losing_screen1.py."),
    ("Lvl 3.py", "Level 3 garden. Door to Garage and vinyl exit door."),
    ("Garage Lvl 3.py", "The garage. Father NPC walks back and forth at y=560, pauses 6 seconds at the car every 2nd lap. AGRO_RADIUS=144px. Katie_Vinyl in a vinyl box. Y-sorting on the car. Lose door to losing_screen1.py."),
    ("Lvl 4.py", "Level 4 garden. Single door to Lvl 4 Gerry room."),
    ("Lvl 4 Gerry room.py", "The final room. VinylPlayer requires all 3 vinyls to activate. After activation the background swaps and a CutsceneBox appears. Pressing E on it launches EndCutScene/ending.py."),
]

for fname, desc in files_info:
    add_h3(fname)
    add_body(desc)

doc.add_page_break()

# ═══ SECTION 15 ═══════════════════════════════════════════════════════════
add_h1("15. Adding a New Level (Template)")
add_body(
    "Copy this template to make a new level. Save it in a new folder (e.g. Level 5/). "
    "Point a Door in the previous level's target_module at this file's path."
)
add_code(
    "import pygame, sys, os\n"
    "from os.path import join\n\n"
    "_HERE = os.path.dirname(os.path.abspath(__file__))\n"
    "_LVL1 = os.path.join(_HERE, '..', 'Level1')\n"
    "if _LVL1 not in sys.path: sys.path.insert(0, _LVL1)\n\n"
    "from hotbar import Overlay, InventoryItem\n"
    "from door   import Door\n"
    "import shared_state\n\n"
    "WINDOW_WIDTH, WINDOW_HEIGHT = 1280, 720\n\n"
    "class Player(pygame.sprite.Sprite):\n"
    "    def __init__(self, groups):\n"
    "        super().__init__(groups)\n"
    "        self.image = pygame.image.load(\n"
    "            r'images\\Player_sprites\\sprite-1-1 (1).png').convert_alpha()\n"
    "        self.rect  = self.image.get_frect(center=(640, 600))\n"
    "        self.direction = pygame.Vector2()\n"
    "        self.speed = 150\n"
    "    def update(self, dt):\n"
    "        keys = pygame.key.get_pressed()\n"
    "        self.direction.x = int(keys[pygame.K_d]) - int(keys[pygame.K_a])\n"
    "        self.direction.y = int(keys[pygame.K_s]) - int(keys[pygame.K_w])\n"
    "        if self.direction.length() > 0:\n"
    "            self.direction = self.direction.normalize()\n"
    "        self.rect.center += self.direction * self.speed * dt\n\n"
    "def run():\n"
    "    pygame.init()\n"
    "    display_surface = pygame.display.set_mode((WINDOW_WIDTH, WINDOW_HEIGHT))\n"
    "    pygame.display.set_caption('My New Level')\n"
    "    clock = pygame.time.Clock()\n"
    "    background = pygame.transform.scale(\n"
    "        pygame.image.load('images/garden.png').convert(),\n"
    "        (WINDOW_WIDTH, WINDOW_HEIGHT))\n"
    "    exit_door = Door(\n"
    "        pos='(640, 650)',\n"
    "        target_module='Level1/Test_level.py',\n"
    "        image_path=None, size=(55, 66))\n"
    "    overlay     = Overlay(Player)\n"
    "    all_sprites = pygame.sprite.Group()\n"
    "    player      = Player(all_sprites)\n"
    "    if getattr(shared_state, 'incoming_hotbar_slots', None):\n"
    "        for i, item in enumerate(shared_state.incoming_hotbar_slots):\n"
    "            overlay.hotbar.slots[i] = item\n"
    "        shared_state.incoming_hotbar_slots = None\n"
    "    running = True\n"
    "    while running:\n"
    "        dt = clock.tick(60) / 1000\n"
    "        for event in pygame.event.get():\n"
    "            if event.type == pygame.QUIT:\n"
    "                pygame.quit(); return\n"
    "            if event.type == pygame.KEYDOWN:\n"
    "                overlay.hotbar.handle_keypress(event)\n"
    "                if event.key == pygame.K_e and exit_door.try_enter(player):\n"
    "                    shared_state.incoming_hotbar_slots = list(overlay.hotbar.slots)\n"
    "                    exit_door.transition(display_surface)\n"
    "                    exit_door.load_next_level(); return\n"
    "        exit_door.update(player)\n"
    "        all_sprites.update(dt)\n"
    "        display_surface.blit(background, (0,0))\n"
    "        all_sprites.draw(display_surface)\n"
    "        exit_door.draw(display_surface)\n"
    "        overlay.display(display_surface)\n"
    "        pygame.display.update()\n"
    "    pygame.quit()\n\n"
    "if __name__ == '__main__':\n"
    "    os.chdir(os.path.join(_HERE, '..'))\n"
    "    run()"
)
doc.add_page_break()

# ═══ SECTION 16 ═══════════════════════════════════════════════════════════
add_h1("16. Common Mistakes and Fixes")

mistakes = [
    ("Player walks through walls",
     "Make sure resolve_collision(player) is called every frame AFTER all_sprites.update(dt), not before."),
    ("Images not found / FileNotFoundError",
     "The working directory must be the project root (the folder containing images/). If running from a level subfolder, add: os.chdir(os.path.join(_HERE, '..'))"),
    ("NPC gets stuck in a corner",
     "Increase FEELER_LEN (try 100) or add more patrol waypoints that steer clear of the problem corner."),
    ("Animation only shows 1-2 frames",
     "Do NOT reset _frame_idx = 0 when direction changes. Let the timer advance it independently."),
    ("Door loads wrong level",
     "The target_module path must be relative to the project root (GerrysHouse/), not the current level folder."),
    ("Items lost between levels",
     "Set shared_state.incoming_hotbar_slots = list(overlay.hotbar.slots) BEFORE calling load_next_level()."),
    ("Window closes immediately on startup",
     "A crash is happening before the game loop starts. Run from terminal: python Level1/Test_level.py to see the error message."),
    ("Y-sorting not working",
     "Use rect.bottom (the feet of the sprite) not rect.centery as the threshold. Compare it against counter_rect.centery."),
    ("Mother/Father doesn't chase",
     "Make sure the _player reference is passed to __init__ and stored as self._player. Check AGRO_RADIUS is large enough."),
]

tbl = doc.add_table(rows=len(mistakes)+1, cols=2)
tbl.style = 'Table Grid'
tbl.rows[0].cells[0].text = "Problem"
tbl.rows[0].cells[1].text = "Fix"
for cell in tbl.rows[0].cells:
    cell.paragraphs[0].runs[0].font.bold = True
for i, (prob, fix) in enumerate(mistakes):
    tbl.rows[i+1].cells[0].text = prob
    tbl.rows[i+1].cells[1].text = fix

doc.add_paragraph("")
add_body("Good luck building your own version of Gerry's House!")

# ── save ──────────────────────────────────────────────────────────────────
out = "How_To_Make_Gerrys_House.docx"
doc.save(out)
print(f"Saved: {out}")
